import os
import logging
from typing import Optional

def extract_text_from_file(file_path: str) -> str:
    """
    Extract text from PDF or DOCX files.
    
    Args:
        file_path (str): Path to the file to extract text from
        
    Returns:
        str: Extracted text content
        
    Raises:
        Exception: If file format is unsupported or extraction fails
    """
    file_extension = os.path.splitext(file_path)[1].lower()
    
    try:
        if file_extension == '.pdf':
            return extract_pdf_text(file_path)
        elif file_extension == '.docx':
            return extract_docx_text(file_path)
        else:
            raise Exception(f"Unsupported file format: {file_extension}")
    except Exception as e:
        logging.error(f"Error extracting text from {file_path}: {str(e)}")
        raise

def extract_pdf_text(file_path: str) -> str:
    """
    Extract text from PDF file using PyMuPDF.
    
    Args:
        file_path (str): Path to PDF file
        
    Returns:
        str: Extracted text content
    """
    try:
        import fitz  # PyMuPDF
        
        doc = fitz.open(file_path)
        text = ""
        
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            text += page.get_text()
            text += "\n"  # Add page break
        
        doc.close()
        return text.strip()
        
    except ImportError:
        # Fallback to pdfplumber if PyMuPDF is not available
        try:
            import pdfplumber
            
            text = ""
            with pdfplumber.open(file_path) as pdf:
                for page in pdf.pages:
                    page_text = page.extract_text()
                    if page_text:
                        text += page_text + "\n"
            
            return text.strip()
            
        except ImportError:
            raise Exception("Neither PyMuPDF nor pdfplumber is available for PDF text extraction. Please install one of these libraries.")
    
    except Exception as e:
        raise Exception(f"Failed to extract text from PDF: {str(e)}")

def extract_docx_text(file_path: str) -> str:
    """
    Extract text from DOCX file using python-docx.
    
    Args:
        file_path (str): Path to DOCX file
        
    Returns:
        str: Extracted text content
    """
    try:
        from docx import Document
        
        doc = Document(file_path)
        text = ""
        
        # Extract text from paragraphs
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    text += cell.text + " "
                text += "\n"
        
        return text.strip()
        
    except ImportError:
        # Fallback to docx2txt if python-docx is not available
        try:
            import docx2txt
            text = docx2txt.process(file_path)
            return text.strip() if text else ""
            
        except ImportError:
            raise Exception("Neither python-docx nor docx2txt is available for DOCX text extraction. Please install one of these libraries.")
    
    except Exception as e:
        raise Exception(f"Failed to extract text from DOCX: {str(e)}")

def clean_extracted_text(text: str) -> str:
    """
    Clean and normalize extracted text.
    
    Args:
        text (str): Raw extracted text
        
    Returns:
        str: Cleaned text
    """
    if not text:
        return ""
    
    # Remove excessive whitespace
    lines = [line.strip() for line in text.split('\n')]
    lines = [line for line in lines if line]  # Remove empty lines
    
    # Join lines with single newlines
    cleaned_text = '\n'.join(lines)
    
    # Remove multiple consecutive spaces
    import re
    cleaned_text = re.sub(r' +', ' ', cleaned_text)
    
    return cleaned_text.strip()
